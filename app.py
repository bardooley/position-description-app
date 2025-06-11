import streamlit as st
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import openai
import requests
from bs4 import BeautifulSoup
import PyPDF2
from collections import Counter
from io import BytesIO

# Set page config
st.set_page_config(
    page_title="Position Description Generator",
    page_icon="📝",
    layout="wide"
)

# Create a container for the header
header_container = st.container()

# Add title image to the right
with header_container:
    col1, col2 = st.columns([3, 1])  # Adjust ratio as needed
    with col2:
        try:
            title_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "title.png")
            st.image(title_path, width=400)
        except Exception as e:
            st.warning(f"Could not load title image: {e}")

# Title and description
st.title("Position Description Generator")
st.markdown("""
This tool helps you generate professional position descriptions for school leadership roles.
Fill out the form below to create your document.
""")

# Create two columns for the form
col1, col2 = st.columns(2)

with col1:
    # School Information
    st.subheader("School Information")
    school_name = st.text_input("School Name")
    position_name = st.text_input("Position Name")
    location = st.text_input("Location")
    website = st.text_input("Website")
    date = st.text_input("Start Date (e.g., July 2026)")
    mission_statement = st.text_area("Mission Statement")

with col2:
    # Consultant Information
    st.subheader("Consultant Information")
    consultant_name = st.text_input("Consultant Name")
    consultant_email = st.text_input("Consultant Email")
    
    # File Uploads
    st.subheader("Required Files")
    notes_file = st.file_uploader("School Notes (PDF)", type=['pdf'])
    logo = st.file_uploader("School Logo", type=['png', 'jpg', 'jpeg'])
    image1 = st.file_uploader("Header Image", type=['png', 'jpg', 'jpeg'])
    image2 = st.file_uploader("Image 2", type=['png', 'jpg', 'jpeg'])
    image3 = st.file_uploader("Image 3", type=['png', 'jpg', 'jpeg'])
    image4 = st.file_uploader("Image 4", type=['png', 'jpg', 'jpeg'])
    image5 = st.file_uploader("Image 5", type=['png', 'jpg', 'jpeg'])


# Preset API keys
openai_api_key = os.environ.get("OPENAI_API_KEY")
google_api_key = os.environ.get("GOOGLE_API_KEY")
google_cse_id = os.environ.get("GOOGLE_SEARCH_ENGINE_ID")

def save_uploaded_file(uploaded_file, filename):
    """Save an uploaded file to the CS&A directory"""
    if uploaded_file is not None:
        try:
            # Create CS&A directory if it doesn't exist
            os.makedirs("CS&A", exist_ok=True)
            
            # Save the file
            file_path = os.path.join("CS&A", filename)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            return file_path
        except Exception as e:
            st.error(f"Error saving file: {str(e)}")
            return None
    return None

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Helvetica")
    rFonts.set(qn("w:hAnsi"), "Helvetica")
    rPr.append(rFonts)
    
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "36")
    rPr.append(sz)
    
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def insert_image_from_memory(doc, image_data, width, after_para=None):
    """Insert an image from memory data after a paragraph"""
    if image_data is not None:
        try:
            if after_para is not None:
                para = doc.insert_paragraph_after(after_para)
            else:
                para = doc.add_paragraph()
            run = para.add_run()
            # Save the image temporarily and use its path
            temp_path = save_uploaded_file(image_data, f"temp_{image_data.name}")
            if temp_path:
                run.add_picture(temp_path, width=width)
                # Clean up the temporary file
                try:
                    os.remove(temp_path)
                except:
                    pass
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return para
        except Exception as e:
            st.error(f"Could not add image: {e}")
    return None

def set_paragraph_format(paragraph):
    """Set consistent paragraph formatting"""
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.right_indent = 0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph

def generate_document():
    """Generate the position description document"""
    try:
        # Read notes content from memory
        notes_content = ""
        if notes_file is not None:
            pdf_reader = PyPDF2.PdfReader(BytesIO(notes_file.getvalue()))
            for page in pdf_reader.pages:
                notes_content += page.extract_text()

        # Set OpenAI API key with error handling
        try:
            client = openai.OpenAI(api_key=openai_api_key)
            # Test the client with a simple request
            client.models.list()
        except Exception as e:
            st.error(f"OpenAI client initialization error: {str(e)}")
            return False, f"Error initializing OpenAI client: {str(e)}", None

        # Generate overview
        overview_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""You are an experienced HR professional crafting the 'Overview' section of a Position Description for a leadership role at an independent school. This is the first section job applicants will see.

Write in a warm, professional tone. The overview should be about 300 words, split into three paragraphs:
1. A brief history of the school and its educational mission
2. A description of the learning environment, programs, and student experience
3. A timely note on leadership, strategy, or facilities that makes this an exciting time to join

Use subtle marketing language—polished but not overblown. Here are two sample overviews to guide your tone and structure:

Example 1:
Founded in 1982, Topeka Collegiate is a PreK–8 independent school that fosters confident, curious learners committed to community and global citizenship. For over 40 years, it has nurtured students through agile, student-centered learning grounded in equity and justice.  
Graduates emerge as critical thinkers and compassionate leaders prepared for lifelong learning. A warm, familial culture supports character development and joyful exploration.  
With the retirement of longtime Head Dr. Lyn Rantz, the school enters a pivotal moment to build on her legacy—including a new academic wing—and launch a strategic plan ahead of its 50th anniversary.

Example 2:
Gulliver Prep, one of the nation's largest independent schools, empowers 2,200 PreK–12 students to pursue excellence through curiosity, integrity, and innovation. Its vibrant Miami campuses blend academic rigor with joyful learning.  
Signature programs in engineering, diplomacy, journalism, and the arts reflect Gulliver's commitment to personalized education. Diversity, global citizenship, and holistic wellness are central to its mission.  
Amid new buildings and deepened partnerships with institutions like Harvard and Stanford, the school seeks a Head of School to help elevate its twin strategic pillars: Next Level Teaching & Learning and Thriving Students."""},
                {"role": "user", "content": f"""Please create the 'Overview' section of a Position Description for the {position_name} position at {school_name} in {location}. Do not output your response with a '**Overview**' first. Base your writing on the following school notes, gathered from a phone call or visit:
{notes_content}
"""}
            ]
        )
        overview = overview_response.choices[0].message.content

        # Generate opportunities and challenges
        opps_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""You are an experienced HR professional drafting the 'Opportunities and Challenges' section of a Position Description for a school leadership role. This section highlights the key strategic priorities, opportunities, and leadership challenges facing the school. It is meant to be clear, forward-looking, and grounded in the school's current context.

Tone: professional, concise, and aspirational.  
Length: About 300 words, ideally formatted as a short contextual paragraph followed by 5–10 bullet points. Make sure all bullet points are in this style: '•' not '-'. Do not output your response with a '**Opportunities and Challenges**' first. 

Common themes to include (only where applicable based on school input):  
• Academic program development  
• Faculty hiring, retention, and growth  
• Enrollment growth and community outreach  
• Strategic planning and governance  
• DEI and culture of belonging  
• Campus improvements  
• Stakeholder engagement and philanthropy

Here are two brief examples for style and structure:

Example 1:
With a committed faculty and strong Board, Topeka Collegiate is well-positioned for its next stage of growth. A recently completed academic wing demonstrates the school's readiness to enhance its programs and deepen its community impact. The next Head will find strong support to take strategic, mission-aligned steps forward.  
Key opportunities and challenges include:  
• Enhancing enrollment through strategic outreach and storytelling  
• Elevating the school's public profile in Topeka and beyond  
• Continuing curricular innovation and experiential learning partnerships  
• Developing a culture of philanthropy to sustain future growth  
• Launching a new strategic planning process with community input

Example 2:
Gulliver Prep offers both scale and flexibility, providing the next Head of School with a robust platform to drive innovation and excellence. The school's diverse community and dynamic programs create space for visionary leadership grounded in collaboration and purpose.  
Key opportunities and challenges include:  
• Advancing a school-wide vision for academic innovation, including AI integration  
• Strengthening professional development pipelines and evaluation systems  
• Expanding off-campus learning: internships, travel, and service  
• Deepening student wellness initiatives and defining what it means to "thrive"  
• Bridging campuses with consistent culture, communication, and leadership"""},
                {"role": "user", "content": f"""Please write the 'Opportunities and Challenges' section for the {position_name} role at {school_name} in {location}, using the following school notes:

{notes_content}
"""}
            ]
        )
        oppChallenges = opps_response.choices[0].message.content

        # Generate qualifications
        quals_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""You are an experienced HR professional writing the 'Qualifications' section of a Position Description for a leadership role at an independent or international school.

This section highlights the qualifications, experience, and personal attributes that the school is seeking in its next leader. It should begin with a short paragraph (2–4 sentences) summarizing the leadership profile the school is looking for, followed by a clean, well-organized bullet-point list of specific qualifications.

Tone: Warm, aspirational, and professional  
Length: ~250–300 words  
Structure: 
1. Opening paragraph about the kind of leader sought  
2. Bullet-point list of qualifications (you may group under subheadings if appropriate). Aim for 12hwh bullet points, and aim for 15 words per bullet point. Every bullet point should be in this style: '•' not '-'.

Here are two brief examples for style and structure:

Example 1:
The next Head of School will be a collaborative, strategic, and inspiring leader who models authenticity, curiosity, and joy. They should be deeply committed to the mission of the school and eager to partner with a talented team to advance its next chapter.

Successful candidates will ideally demonstrate:
• Leadership experience in an independent or international school  
• Deep understanding of curriculum development and faculty support  
• Strong communication skills and cultural competency  
• Strategic thinking paired with operational effectiveness  
• Warmth, humility, humor, and a student-centered mindset  
• A master's degree in education or related field

Example 2:
Gulliver seeks a visible, emotionally intelligent leader with a bold vision for teaching, learning, and community life. The ideal candidate will be energized by the school's diversity, scale, and ambition—and grounded in inclusive, servant-minded leadership.

Ideal qualifications include:
• Experience leading in large, complex K–12 environments  
• Success recruiting and retaining diverse, mission-aligned teams  
• A collaborative, transparent leadership style  
• Strength in governance, strategy, and data-driven decision-making  
• Courage to make difficult decisions with integrity  
• Advanced degree and proven commitment to professional growth

Do not output your response with a '**Qualifications**' first. Do not break your bullet points up into sections of bullet points. Output all bullet points in order. For example, do not break the bullet points up into sections like 'Educational and Leadership Experience.'"""},
                {"role": "user", "content": f"""Please write the 'Qualifications' section of a Position Description for the {position_name} role at {school_name} in {location}, using the following school notes:

{notes_content}
"""}
            ]
        )
        qualifications = quals_response.choices[0].message.content

        # Google Custom Search API for relevant websites
        def google_search(query, api_key, cse_id):
            url = "https://www.googleapis.com/customsearch/v1"
            params = {
                "q": query,
                "key": api_key,
                "cx": cse_id,
                "num": 1
            }
            response = requests.get(url, params=params)
            results = response.json()
            if "items" in results and len(results["items"]) > 0:
                return results["items"][0]["link"]
            else:
                return None

        # Get relevant websites
        websites_dict = {
            "School Website": google_search(f"{school_name} official site", google_api_key, google_cse_id),
            "Town Information": google_search(f"Town of {location} official site", google_api_key, google_cse_id),
            "School History": google_search(f"{school_name} history", google_api_key, google_cse_id)
        }

        # Create document
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        # Enable different first page header/footer only for first section
        section.different_first_page_header_footer = True
        
        # Add footer to first page
        try:
            # Get the directory where the script is located
            script_dir = os.path.dirname(os.path.abspath(__file__))
            footer_path = os.path.join(script_dir, "footer.png")
            with open(footer_path, 'rb') as f:
                footer_data = f.read()
            # Get the first page footer section
            footer = section.first_page_footer
            # Add the footer image
            footer_para = footer.paragraphs[0]
            footer_run = footer_para.add_run()
            footer_run.add_picture(BytesIO(footer_data), width=Inches(7.2))
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            st.warning(f"Could not add footer: {e}")

        # Add 'hello' to footer of all other pages
        try:
            regular_footer = section.footer
            # Left-aligned text
            footer_para = regular_footer.paragraphs[0]
            footer_run = footer_para.add_run('The Search Group | Carney, Sandoe & Associates')
            footer_run.font.size = Pt(12)
            footer_run.font.name = 'Helvetica'
            footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Right-aligned page number
            right_para = regular_footer.add_paragraph()
            right_run = right_para.add_run('Page ')
            right_run.font.size = Pt(12)
            right_run.font.name = 'Helvetica'
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            right_run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            right_run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            right_run._r.append(fldChar2)
        except Exception as e:
            st.warning(f"Could not add regular footer: {e}")

        # Add header table
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(5.5)
        header_table.columns[1].width = Inches(0.9)

        # Add header content
        cell_left = header_table.cell(0, 0)
        title_para = cell_left.paragraphs[0]
        title_run = title_para.add_run(f"{position_name.upper()} SEARCH")
        title_run.bold = True
        title_run.font.size = Pt(23)
        title_run.font.name = "Helvetica"
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add school name
        school_para = cell_left.add_paragraph()
        school_run = school_para.add_run(school_name.upper())
        school_run.bold = True
        school_run.font.size = Pt(16)
        school_run.font.name = "Helvetica"
        school_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add location, website, and start date
        for text in [location, website, f"Start Date: {date}"]:
            para = cell_left.add_paragraph(text)
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.name = "Helvetica"
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add logo from memory
        if logo is not None:
            cell_right = header_table.cell(0, 1)
            logo_para = cell_right.paragraphs[0]
            logo_run = logo_para.add_run()
            logo_run.add_picture(BytesIO(logo.getvalue()), width=Inches(2.0))
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Extract dominant color from logo
            try:
                with PILImage.open(BytesIO(logo.getvalue())) as img:
                    img = img.convert('RGB')
                    img = img.resize((50, 50))
                    pixels = list(img.getdata())
                    most_common_colors = Counter(pixels).most_common(5)
                    
                    def is_white_grey_black(rgb):
                        r, g, b = rgb
                        if r > 240 and g > 240 and b > 240:  # White
                            return True
                        if r < 20 and g < 20 and b < 20:  # Black
                            return True
                        if abs(r-g) < 10 and abs(r-b) < 10 and abs(g-b) < 10 and 60 < r < 220:  # Grey
                            return True
                        return False
                    
                    for color, _ in most_common_colors:
                        if not is_white_grey_black(color):
                            mission_title_color = color
                            break
                    else:
                        mission_title_color = most_common_colors[0][0]
            except Exception as e:
                st.warning(f"Could not extract color from logo: {e}")
                mission_title_color = (0, 0, 0)  # Default to black
        else:
            mission_title_color = (0, 0, 0)  # Default to black

        # Add image1 from memory
        if image1 is not None:
            image_table = doc.add_table(rows=1, cols=1)
            image_table.autofit = False
            image_table.columns[0].width = Inches(6.5)
            image_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = image_table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(BytesIO(image1.getvalue()), width=Inches(6.5))
            
            # Create new section without different first page
            section = doc.add_section()
            section.start_new_page = True
            section.different_first_page_header_footer = False
            
            # Link to previous section's footer
            section.footer.is_linked_to_previous = True

        # Add mission statement
        mission_title_para = doc.add_paragraph()
        mission_title_run = mission_title_para.add_run('Mission Statement')
        mission_title_run.bold = True
        mission_title_run.font.size = Pt(24)
        mission_title_run.font.name = 'Helvetica'
        mission_title_run.font.color.rgb = RGBColor(*mission_title_color)
        mission_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mission_title_para.paragraph_format.space_before = Pt(0)
        mission_title_para.paragraph_format.space_after = Pt(12)
        mission_title_para.paragraph_format.page_break_before = True

        # Add mission statement text
        mission_text_para = doc.add_paragraph()
        mission_text_run = mission_text_para.add_run(mission_statement)
        mission_text_run.font.size = Pt(13)
        mission_text_run.font.name = 'Helvetica'
        mission_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mission_text_para.paragraph_format.line_spacing = 0.99

        # Insert image2 from memory
        if image2 is not None:
            insert_image_from_memory(doc, image2, Inches(7))

        # Add overview
        overview_title_para = doc.add_paragraph()
        overview_title_run = overview_title_para.add_run('Overview')
        overview_title_run.bold = True
        overview_title_run.font.size = Pt(24)
        overview_title_run.font.name = 'Helvetica'
        overview_title_run.font.color.rgb = RGBColor(*mission_title_color)
        overview_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add overview text
        overview_text_para = doc.add_paragraph()
        overview_text_run = overview_text_para.add_run(overview)
        overview_text_run.font.size = Pt(13)
        overview_text_run.font.name = 'Helvetica'
        overview_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        overview_text_para.paragraph_format.line_spacing = 0.99

        # Insert image3 from memory
        if image3 is not None:
            insert_image_from_memory(doc, image3, Inches(7))

        # Add opportunities and challenges
        opps_title_para = doc.add_paragraph()
        opps_title_run = opps_title_para.add_run('Opportunities and Challenges')
        opps_title_run.bold = True
        opps_title_run.font.size = Pt(24)
        opps_title_run.font.name = 'Helvetica'
        opps_title_run.font.color.rgb = RGBColor(*mission_title_color)
        opps_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Process opportunities and challenges text
        phrase = "Key opportunities and challenges include:"
        if phrase in oppChallenges:
            before, after = oppChallenges.split(phrase, 1)
            opps_text_para = doc.add_paragraph()
            if before.strip():
                run_before = opps_text_para.add_run(before)
                run_before.font.size = Pt(13)
                run_before.font.name = 'Helvetica'
            run_bold = opps_text_para.add_run(phrase)
            run_bold.bold = True
            run_bold.font.size = Pt(13)
            run_bold.font.name = 'Helvetica'
            opps_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            opps_text_para.paragraph_format.line_spacing = 0.99

            for line in after.split('\n'):
                line = line.strip()
                if line.startswith('•'):
                    bullet_para = doc.add_paragraph()
                    bullet_run = bullet_para.add_run(line)
                    bullet_run.font.size = Pt(13)
                    bullet_run.font.name = 'Helvetica'
                    bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    bullet_para.paragraph_format.left_indent = Pt(9)
                    bullet_para.paragraph_format.first_line_indent = Pt(-9)
                    bullet_para.paragraph_format.space_after = Pt(1)
                    bullet_para.paragraph_format.line_spacing = 0.99
                elif line:
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.size = Pt(13)
                    run.font.name = 'Helvetica'
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 0.99
        else:
            opps_text_para = doc.add_paragraph()
            opps_text_run = opps_text_para.add_run(oppChallenges)
            opps_text_run.font.size = Pt(13)
            opps_text_run.font.name = 'Helvetica'
            opps_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            opps_text_para.paragraph_format.line_spacing = 0.99

        # Insert image4 from memory
        if image4 is not None:
            insert_image_from_memory(doc, image4, Inches(7))

        # Add qualifications
        quals_title_para = doc.add_paragraph()
        quals_title_run = quals_title_para.add_run('Qualifications and Personal Attributes')
        quals_title_run.bold = True
        quals_title_run.font.size = Pt(24)
        quals_title_run.font.name = 'Helvetica'
        quals_title_run.font.color.rgb = RGBColor(*mission_title_color)
        quals_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Process qualifications text
        if '•' in qualifications:
            for line in qualifications.split('\n'):
                line = line.strip()
                if line.startswith('•'):
                    bullet_para = doc.add_paragraph()
                    bullet_run = bullet_para.add_run(line)
                    bullet_run.font.size = Pt(13)
                    bullet_run.font.name = 'Helvetica'
                    bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    bullet_para.paragraph_format.left_indent = Pt(9)
                    bullet_para.paragraph_format.first_line_indent = Pt(-9)
                    bullet_para.paragraph_format.space_after = Pt(1)
                    bullet_para.paragraph_format.line_spacing = 0.99
                elif line:
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.size = Pt(13)
                    run.font.name = 'Helvetica'
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 0.99
        else:
            quals_text_para = doc.add_paragraph()
            quals_text_run = quals_text_para.add_run(qualifications)
            quals_text_run.font.size = Pt(13)
            quals_text_run.font.name = 'Helvetica'
            quals_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            quals_text_para.paragraph_format.line_spacing = 0.99

        # Insert image5 from memory
        if image5 is not None:
            insert_image_from_memory(doc, image5, Inches(7))

        # Add learn more section
        learn_title_para = doc.add_paragraph()
        learn_title_run = learn_title_para.add_run('Learn More')
        learn_title_run.bold = True
        learn_title_run.font.size = Pt(24)
        learn_title_run.font.name = 'Helvetica'
        learn_title_run.font.color.rgb = RGBColor(*mission_title_color)
        learn_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add learn more subtitle
        learn_subtitle_para = doc.add_paragraph()
        learn_subtitle_run = learn_subtitle_para.add_run(f'Click on the links below to learn more about {school_name}.')
        learn_subtitle_run.font.size = Pt(13)
        learn_subtitle_run.font.name = 'Helvetica'
        learn_subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add links table
        num_links = len(websites_dict)
        num_rows = (num_links + 1) // 2
        links_table = doc.add_table(rows=num_rows, cols=2)
        links_table.autofit = False
        links_table.columns[0].width = Inches(3.25)
        links_table.columns[1].width = Inches(3.25)

        websites_list = list(websites_dict.items())
        for i, (website_name, url) in enumerate(websites_list):
            row = i // 2
            col = i % 2
            cell = links_table.cell(row, col)
            para = cell.paragraphs[0]
            add_hyperlink(para, url, website_name)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add to apply section
        apply_title_para = doc.add_paragraph()
        apply_title_para.paragraph_format.space_before = Pt(24)
        apply_title_run = apply_title_para.add_run('To Apply')
        apply_title_run.bold = True
        apply_title_run.font.size = Pt(24)
        apply_title_run.font.name = 'Helvetica'
        apply_title_run.font.color.rgb = RGBColor(*mission_title_color)
        apply_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add application instructions
        apply_text_para = doc.add_paragraph()
        apply_text_run = apply_text_para.add_run("Interested and qualified candidates are invited to contact the consultants in confidence. Candidates will ultimately need to submit the following materials as separate PDF documents:")
        apply_text_run.font.size = Pt(13)
        apply_text_run.font.name = 'Helvetica'
        apply_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_text_para.paragraph_format.line_spacing = 0.99

        # Add bullet points
        apply_bullets = [
            "A cover letter expressing their interest in this particular position;",
            "A current résumé;",
            "A statement of educational and leadership philosophy;",
            "A list of five professional references with name, relationship, phone number, and email address of each (references will not be contacted without the candidate's permission) to:"
        ]
        for bullet in apply_bullets:
            bullet_para = doc.add_paragraph()
            bullet_run = bullet_para.add_run(f"• {bullet}")
            bullet_run.font.size = Pt(13)
            bullet_run.font.name = 'Helvetica'
            bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bullet_para.paragraph_format.left_indent = Pt(18)
            bullet_para.paragraph_format.space_after = Pt(1)
            bullet_para.paragraph_format.line_spacing = 0.99

        # Add consultant information
        consultant_name_para = doc.add_paragraph()
        consultant_name_para.paragraph_format.space_before = Pt(24)
        consultant_name_run = consultant_name_para.add_run(consultant_name)
        consultant_name_run.bold = True
        consultant_name_run.font.size = Pt(13)
        consultant_name_run.font.name = 'Helvetica'
        consultant_name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        consultant_name_para.paragraph_format.space_after = Pt(1)

        consultant_title_para = doc.add_paragraph()
        consultant_title_run = consultant_title_para.add_run("Consultant")
        consultant_title_run.font.size = Pt(13)
        consultant_title_run.font.name = 'Helvetica'
        consultant_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        consultant_title_para.paragraph_format.space_after = Pt(1)

        consultant_email_para = doc.add_paragraph()
        consultant_email_run = consultant_email_para.add_run(consultant_email)
        consultant_email_run.font.size = Pt(13)
        consultant_email_run.font.name = 'Helvetica'
        consultant_email_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        consultant_email_para.paragraph_format.space_after = Pt(1)

        # Save document to memory
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return True, "Document generated successfully!", docx_buffer

    except Exception as e:
        return False, f"Error generating document: {str(e)}", None

# Generate button
if st.button("Generate Position Description"):
    if not all([school_name, position_name, location, website, date, mission_statement, 
                consultant_name, consultant_email]):
        st.error("Please fill in all required fields.")
    elif not openai_api_key or not google_api_key or not google_cse_id:
        st.error("API keys not found in environment variables. Please ensure OPENAI_API_KEY, GOOGLE_API_KEY, and GOOGLE_SEARCH_ENGINE_ID are set.")
    else:
        with st.spinner("Generating document..."):
            success, message, docx_buffer = generate_document()
            if success:
                st.success(message)
                # Provide download link for DOCX only
                st.download_button(
                    label="Download Word Document",
                    data=docx_buffer.getvalue(),
                    file_name="position_description.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(message) 